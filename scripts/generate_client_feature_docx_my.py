#!/usr/bin/env python3
"""Generate Learn2Success client feature guide — Myanmar (.docx)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Learn2Success-Client-Feature-Guide-Myanmar.docx"

FONT = "Pyidaungsu"
FONT_FALLBACK = "Myanmar Text"

GOLD = RGBColor(0x78, 0x59, 0x00)
GOLD_DARK = RGBColor(0x6D, 0x51, 0x00)
SKY = RGBColor(0x00, 0x63, 0x99)
GREEN = RGBColor(0x00, 0x6E, 0x1C)
SLATE = RGBColor(0x47, 0x55, 0x69)
BODY = RGBColor(0x23, 0x2C, 0x36)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FILL_GOLD = "FFF9EB"
FILL_SKY = "E3F2FD"
FILL_GREEN = "F3FFF2"
FILL_HEADER = "785900"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def _apply_font(run, size=11, bold=False, color=BODY, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_FALLBACK)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_run(paragraph, text, *, size=11, bold=False, color=BODY, italic=False):
    run = paragraph.add_run(text)
    _apply_font(run, size, bold, color, italic)
    return run


def style_cell_fonts(cell, size=10, bold=False, color=BODY):
    for p in cell.paragraphs:
        for run in p.runs:
            _apply_font(run, size, bold, color)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p, FILL_SKY)
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Learn2Success", size=34, bold=True, color=GOLD)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p2, FILL_SKY)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, "လုပ်ဆောင်ချက် လမ်းညွှန်ချက် (မြန်မာ)", size=18, bold=True, color=SKY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p3, FILL_SKY)
    p3.paragraph_format.space_after = Pt(24)
    add_run(p3, "အတူတကွ သင်ယူကြမယ်၊ အတူတကွ အောင်မြင်ကြမယ်။", size=12, italic=True, color=GOLD_DARK)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(24)
    add_run(
        meta,
        f"ပေးအပ်သည့်ရက် — {date.today().strftime('%Y-%m-%d')}  ·  မိုဘိုင်းနှင့်ဒက်စ်တော့ နှစ်မျိုးလုံးအသုံးပြုနိုင်သော ဝဘ်အက်ပ်",
        size=10,
        color=SLATE,
    )


def add_section_heading(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, title, size=16, bold=True, color=GOLD)

    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(10)
        add_run(s, subtitle, size=10, color=SLATE)


def add_subheading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, size=12, bold=True, color=SKY)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    add_run(p, text, size=11, color=BODY)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.left_indent = Inches(0.25)
        add_run(p, item, size=11, color=BODY)


def add_feature_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "လုပ်ဆောင်ချက်"
    headers[1].text = "ရရှိမည့်အကျိုးကျေးဇူး"
    for cell in headers:
        set_cell_shading(cell, FILL_HEADER)
        style_cell_fonts(cell, bold=True, color=WHITE)

    for i, (feature, detail) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        c0.text = feature
        c1.text = detail
        fill = FILL_GOLD if i % 2 == 1 else "FFFFFF"
        set_cell_shading(c0, fill)
        set_cell_shading(c1, fill)
        style_cell_fonts(c0)
        style_cell_fonts(c1)

    doc.add_paragraph()


def add_role_matrix(doc: Document) -> None:
    headers = ["လုပ်ဆောင်ချက်", "ဧည့်သည်", "ကျောင်းသား", "ဆရာ", "စီမံခန့်ခွဲသူ"]
    data = [
        ("မိတ်ဆက်စာမျက်နှာ နှင့် ကုန်အမှတ်တံဆိပ်", "ရှိ", "—", "—", "—"),
        ("ဝင်ရောက် / ထွက်ရန်", "ရှိ", "ရှိ", "ရှိ", "ရှိ"),
        ("အပတ်စဉ် သင်ယူမှုလမ်းကြောင်း", "—", "ရှိ", "—", "—"),
        ("ကွစ်ဇ်ဖြေဆိုခြင်း (ရွေးချယ် / ပြောပြ)", "—", "ရှိ", "—", "—"),
        ("အလိုအလျောက် အမှတ်ပေး (ရွေးချယ်မှုတစ်မျိုးတည်း)", "—", "ရှိ", "—", "—"),
        ("ဆက်လက်ဖြေ / ပြန်လည်ကြိုးစား", "—", "ရှိ", "—", "—"),
        ("ကိုယ်ပိုင် အမှတ်နှင့် မှတ်တမ်း", "—", "ရှိ", "—", "—"),
        ("ဒက်ရှ်ဘုတ် နှင့် ဇယားများ", "—", "—", "ရှိ", "ရှိ"),
        ("အသံဖြေဆိုမှု အမှတ်ပေးခြင်း", "—", "—", "ရှိ", "—"),
        ("အသုံးပြုသူ စီမံခန့်ခွဲမှု", "—", "—", "—", "ရှိ"),
        ("ကွစ်ဇ်ရေးဆွဲခြင်း", "—", "—", "—", "ရှိ"),
        ("အပတ်နှင့် ကွစ်ဇ်အစဉ် စံခြင်း", "—", "—", "—", "ရှိ"),
    ]

    table = doc.add_table(rows=1 + len(data), cols=5)
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        style_cell_fonts(cell, size=9, bold=True, color=WHITE)

    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            fill = FILL_SKY if i % 2 == 0 else "FFFFFF"
            if j == 0:
                fill = FILL_GOLD if i % 2 == 1 else "FFFFFF"
            set_cell_shading(cell, fill)
            color = GREEN if val == "ရှိ" else BODY
            style_cell_fonts(cell, size=9, bold=(val == "ရှိ"), color=color)

    doc.add_paragraph()


def add_demo_accounts(doc: Document) -> None:
    accounts = [
        ("student", "password", "အပတ် ၁ သာ ဖွင့်ထား — နောက်အပတ်များ သော့ခတ်ထားသည်ကို ပြသရန်"),
        ("student_allweeks", "password", "အပတ် ၄ ခုလုံး ဖွင့်ထား — လမ်းကြောင်း အပြည့်အစုံ"),
        ("student_jan1", "password", "၂၀၂၆ ဇန်နဝါရီ ၁ ရက် မှတ်ပုံတင်ထား — ရက်စွဲအလိုက် ဖွင့်မှု"),
        ("teacher", "password", "အမှတ်ပေးဒက်ရှ်ဘုတ် နှင့် စာရင်း"),
        ("admin", "password", "အသုံးပြုသူ၊ ကွစ်ဇ်၊ အပတ် စီမံခန့်ခွဲမှု အပြည့်"),
    ]

    table = doc.add_table(rows=1 + len(accounts), cols=3)
    table.style = "Table Grid"

    for j, h in enumerate(["အသုံးပြုသူအမည်", "စကားဝှက်", "သရုပ်ပြရန် အသုံးပြုရန်"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        style_cell_fonts(cell, bold=True, color=WHITE)

    for i, (user, pwd, note) in enumerate(accounts, start=1):
        row = table.rows[i].cells
        row[0].text = user
        row[1].text = pwd
        row[2].text = note
        fill = FILL_GREEN if i % 2 == 1 else "FFFFFF"
        for cell in row:
            set_cell_shading(cell, fill)
            style_cell_fonts(cell)

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title_block(doc)

    add_section_heading(
        doc,
        "၁။ ပလက်ဖောင်း အကျဉ်းချုပ်",
        "အင်္ဂလိပ်စကား လေ့ကျင့်ခြင်းနှင့် အကဲဖြတ်ခြင်းအတွက် ဝဘ်အက်ပ်တစ်ခု။",
    )
    add_body(
        doc,
        "Learn2Success သည် ကျောင်းသားများအား စနစ်တကျ ကွစ်ဇ်များဖြင့် အင်္ဂလိပ်စကား လေ့ကျင့်ခွင့် "
        "ပေးပါသည် — အထူးသဖြင့် ဘရောက်ဇာတွင် အသံဖမ်းပြီး ပြောကြားသော အဖြေများ။ ဆရာများက အမှတ်နှင့် "
        "တုံ့ပြန်ချက်ဖြင့် အမှတ်ပေးပါသည်။ စီမံခန့်ခွဲသူများက အသုံးပြုသူများ၊ ကွစ်ဇ်အကြောင်းအရာနှင့် "
        "အချိန်အလိုက် ဖွင့်လှစ်သော အပတ်အလိုက် သင်ယူမှုလမ်းကြောင်းကို စီမံပါသည်။",
    )
    add_body(
        doc,
        "ဤထုတ်ကုန်သည် မိုဘိုင်းဖုန်းနှင့် ကွန်ပျူတာ နှစ်မျိုးလုံးတွင် အသုံးပြုနိုင်သော ဝဘ်ဆိုက်ဖြစ်ပြီး "
        "အက်ပ်စတိုးမှ သီးခြားဒေါင်းလုဒ် လုပ်ရန် မလိုပါ။",
    )

    add_section_heading(doc, "၂။ အသုံးပြုသူ အမျိုးအစားများ")
    add_role_matrix(doc)

    add_section_heading(
        doc,
        "၃။ ဧည့်သည် — မိတ်ဆက်စာမျက်နှာ နှင့် ဝင်ရောက်ခြင်း",
        "လော့ဂ်အင်မဝင်မီ အများပြည်သူကြည့်နိုင်သော စာမျက်နှာများ။",
    )
    add_feature_table(
        doc,
        [
            (
                "မိတ်ဆက်စာမျက်နှာ",
                "လိုဂို၊ စာတန်း၊ ပုံလှည့်ပြသမှု (ဘယ်ညာ ခလုတ်နှင့် လက်ဖြင့် ပွတ်ဆွဲခြင်း)၊ "
                "နားထောင် / ပြောပြ / အောင်မြင်ပါစေ ဟူသော အဓိကအချက်များ ပါဝင်သည်။",
            ),
            (
                "ဝင်ရောက်ခြင်း",
                "အသုံးပြုသူအမည် နှင့် စကားဝှက်ဖြင့် ဝင်ရောက်နိုင်ပြီး မှတ်မိထားရန် ရွေးချယ်နိုင်သည်။",
            ),
            (
                "အခန်းကဏ္ဍအလိုက် လမ်းညွှန်",
                "ဝင်ပြီးနောက် ကျောင်းသား လမ်းကြောင်း၊ ဆရာ ဒက်ရှ်ဘုတ် သို့မဟုတ် စီမံခန့်ခွဲသူ ဒက်ရှ်ဘုတ်သို့ "
                "သက်ဆိုင်ရာအတိုင်း ရောက်ရှိသည်။",
            ),
        ],
    )

    add_section_heading(
        doc,
        "၄။ ကျောင်းသား အတွေ့အကြုံ",
        "မိုဘိုင်းအတွက် အသင့်လျက်ရှိသော သင်ယူမှုလမ်းကြောင်း နှင့် ကွစ်ဇ်ဖြေဆိုမှု။",
    )
    add_subheading(doc, "၄.၁ အပတ်အလိုက် လမ်းကြောင်း")
    add_bullets(
        doc,
        [
            "ကျောင်းသား ပင်မစာမျက်နှာတွင် ဂိမ်းလမ်းကြောင်းပုံစံ အပတ်အလိုက် လမ်းကြောင်း ပြသသည်။",
            "ကျောင်းသား မှတ်ပုံတင်သည့်ရက်မှ အပတ်အလိုက် ဖွင့်လှစ်သည် (အပတ် ၁ — ရက် ၀၊ အပတ် ၂ — ရက် ၇ ...)။",
            "သော့ခတ်ထားသော အပတ်များကို မြင်ရသော်လည်း မဖွင့်မီ ဝင်ရောက်၍ မရပါ။",
            "အပတ်တစ်ခုအတွင်း ကွစ်ဇ်အဆင့်များ — ရွေးချယ် / ပြောပြ အမျိုးအစားနှင့် အချိန်ကန့်သတ်မှု ပြသသည်။",
        ],
    )
    add_subheading(doc, "၄.၂ ကွစ်ဇ်နှင့် ကြိုးစားမှုများ")
    add_bullets(
        doc,
        [
            "ကွစ်ဇ်အသေးစိတ်တွင် ဖော်ပြချက်၊ အမှတ်စုစုပေါင်း၊ မေးခွန်းအရေအတွက်၊ ကြိုးစားမှုမှတ်တမ်း နှင့် နောက်ဆုံးအခြေအနေ ပါဝင်သည်။",
            "အခြေအနေများ — ဆက်လက်ဖြေဆိုနေဆဲ၊ ဆရာအမှတ်ပေးရန် စောင့်ဆိုင်း၊ အမှတ်ပေးပြီး။",
            "စတင်ပါ၊ ဆက်လက်ဖြေပါ၊ ပြန်လည်ကြိုးစားပါ — ကွစ်ဇ်တစ်ခုလျှင် ကြိုးစားမှုအကြိမ်ရေ မကန့်သတ်ပါ။",
            "ကွစ်ဇ်တစ်ခုလျှင် ဆက်လက်ဖြေဆိုနေဆဲ ကြိုးစားမှု တစ်ခုသာ ရှိနိုင်သည်။",
        ],
    )
    add_subheading(doc, "၄.၃ ကွစ်ဇ်ဖြေဆိုခြင်း")
    add_bullets(
        doc,
        [
            "ရွေးချယ်ပါ (အများရွေးတစ်ခု) — တစ်ခုရွေးချယ်သည်။ ကွစ်ဇ်တစ်ခုလုံး ရွေးချယ်မှုတည်းဆိုလျှင် အလိုအလျောက် အမှတ်ပေး။",
            "ပြောပါ (အသံဖမ်းခြင်း) — နားထောင်ပါ (စာသားမှ အသံ)、မှတ်တမ်းတင်၊ နားထောင်ပြန်ကြည့်、ပြန်လည်မှတ်တမ်းတင်။",
            "ကွစ်ဇ်အချိန်ကန့်သတ်မှု ရှိနိုင်ပြီး အချိန်ကုန်လျှင် အလိုအလျောက် တင်သွင်းသည်။",
            "မေးခန်းတိုင်းတွင် တိုးတက်မှုပြသသည်။ ကျောင်းသား ရှေ့သို့သွားသည့်အခါ အဖြေများ သိမ်းဆည်းသည်။",
        ],
    )

    add_section_heading(
        doc,
        "၅။ ဆရာ အတွေ့အကြုံ",
        "အကြောင်းအရာ တည်းဖြတ်ခြင်းမဟုတ် — အမှတ်ပေးခြင်းကို အဓိကထားသည်။",
    )
    add_feature_table(
        doc,
        [
            ("ဒက်ရှ်ဘုတ်", "အချက်အလက်ကတ်များ၊ ဇယားများ၊ ဦးစားပေးစာရင်း (အဟောင်းဆုံး ဦးစွာ) နှင့် အမြန်လင့်ခ်များ။"),
            ("အမှတ်ပေးစာရင်း", "စစ်ဆေးရန်အဆင်သင့်၊ ဆက်လက်ဖြေဆိုနေဆဲ၊ အမှတ်ပေးပြီး၊ အားလုံး — စာမျက်နှာခွဲထားသည်။"),
            ("ကျောင်းသားအလိုက်", "ကျောင်းသားတစ်ဦးချင်းစီ၏ ကြိုးစားမှုများကို အုပ်စုဖွဲ့ပြသသည်။"),
            ("အမှတ်ပေးစာမျက်နှာ", "ကွစ်ဇ်အချက်အလက် ဖတ်ရှုရုံသာ။ အသံဖြေများကို နားထောင်နိုင်သည်။ မေးခွန်းတိုင်း အမှတ်နှင့် တုံ့ပြန်ချက်။"),
            ("အမှတ်စစ်ဆေးမှု", "စုစုပေါင်းအမှတ် ဘားတန်း။ ကွစ်ဇ်အများဆုံးအမှက် နှင့် မေးခွန်းတိုင်း အများဆုံးထက် ကျော်လွန်၍ မရ။ မမှန်လျှင် သိမ်းဆည်း ခလုတ် ပိတ်ထားသည်။"),
        ],
    )

    add_section_heading(
        doc,
        "၆။ စီမံခန့်ခွဲသူ အတွေ့အကြုံ",
        "ပရိုဂရမ် နှင့် အကောင့် စီမံခန့်ခွဲမှု အပြည့်အစုံ။",
    )
    add_feature_table(
        doc,
        [
            ("ဒက်ရှ်ဘုတ်", "ကျောင်းသား/ဆရာအရေအတွက်၊ တက်ကြွသော ကွစ်ဇ်များ၊ အမှတ်ပေးရန်စောင့်ဆိုင်း၊ ဇယားများ၊ မကြာသေးမီ လုပ်ဆောင်ချက်များ။"),
            ("အသုံးပြုသူများ", "ကျောင်းသား၊ ဆရာ၊ စီမံခန့်ခွဲသူ ဖန်တီး/တည်းဖြတ်/ဖျက်ခြင်း။ စကားဝှက် စည်းမျဉ်းများ။ ကိုယ်ပိုင် admin အကောင့် ဖျက်၍ မရ။"),
            ("ကွစ်ဇ်များ", "မေးခွန်း ၁ မှ ၁၀၀ အထိ။ အမျိုးအစား နှစ်မျိုး — ရွေးချယ်ပါ နှင့် ပြောပါ။ ကွစ်ဇ်တိုင်း အချိန်ကန့်သတ် ရွေးချယ်နိုင်သည်။"),
            ("အပတ်များ", "အပတ်ဖန်တီး/တည်းဖြတ်။ ကွစ်ဇ်ထည့်သွင်း၊ ဆွဲယူ၍ အစဉ်လိုက်၊ အပတ်မှ ဖယ်ရှား (ကွစ်ဇ်ကိုယ်တိုင် မဖျက်ပါ)။"),
            ("ကွစ်ဇ် ↔ အပတ် စည်းမျဉ်း", "အပတ်နှင့် ချိတ်ဆက်မှုကို အပတ် တည်းဖြတ်စာမျက်နှာတွင်သာ လုပ်သည် — ကွစ်ဇ်ဖောင်တွင် မပါ။"),
        ],
    )

    add_section_heading(doc, "၇။ မေးခွန်း အမျိုးအစားများ")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for j, h in enumerate(["အမျိုးအစား", "ကျောင်းသား မြင်ရမည့်အရာ", "အမှတ်ပေးမှု"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        style_cell_fonts(cell, bold=True, color=WHITE)
    rows = [
        ("ရွေးချယ်ပါ", "အများရွေးတစ်ခု — တစ်ခုရွေးချယ်ပါ", "ကွစ်ဇ်တစ်ခုလုံး ရွေးချယ်မှုတည်းဆိုလျှင် အလိုအလျောက်"),
        ("ပြောပါ", "ဘရောက်ဇာတွင် အသံဖမ်းပြီး ပြောကြားပါ", "ဆရာက အမှတ်နှင့် တုံ့ပြန်ချက် ပေးသည်"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            set_cell_shading(cell, FILL_GOLD if i == 1 else FILL_SKY)
            style_cell_fonts(cell)
    doc.add_paragraph()

    add_section_heading(doc, "၈။ အရည်အသွေး နှင့် လုံခြုံရေး")
    add_bullets(
        doc,
        [
            "အခန်းကဏ္ဍ သုံးမျိုး — ကျောင်းသားသည် ဝန်ထမ်းနယ်ပယ်သို့ ဝင်ရောက်၍ မရပါ။",
            "ကျောင်းသားသည် ကိုယ်ပိုင် တင်သွင်းမှုများကိုသာ ကြည့်နိုင်သည်။",
            "ပိတ်ထားသော သို့မဟုတ် သော့ခတ်ထားသော ကွစ်ဇ်များကို ဝင်ရောက်ခွင့် ပိတ်ထားသည်။",
            "ဝင်ရောက်၊ လမ်းကြောင်း၊ အမှတ်ပေး၊ စီမံခန့်ခွဲမှု၊ ဒက်ရှ်ဘုတ် အပါအဝင် အလိုအလျောက် စမ်းသပ်မှု ၇၃ ခု။",
            "အသံဖိုင်များကို ဆာဗာတွင် သိမ်းဆည်းသည်။ မေးခွန်းတစ်ခုလျှင် အများဆုံး ၃၀ စက္ကန့် မှတ်တမ်းတင် UI။",
        ],
    )

    add_section_heading(doc, "၉။ သရုပ်ပြအတွက် အကောင့်များ")
    add_demo_accounts(doc)

    add_section_heading(
        doc,
        "၁၀။ ဤထုတ်ဝေမှုတွင် မပါဝင်သေးသော အရာများ",
        "အနာဂတ် တိုးချဲ့မှုအတွက် ဖြစ်နိုင်သော အကြံပြုချက်များ — လက်ရှိပေးအပ်မှုတွင် မပါဝင်ပါ။",
    )
    add_bullets(
        doc,
        [
            "iOS / Android မူရင်းအက်ပ် (ယနေ့ ဝဘ်သာ)။",
            "ကျောင်းသား ကိုယ်တိုင် မှတ်ပုံတင်ခြင်း သို့မဟုတ် စကားဝှက် ပြန်လည်သတ်မှတ်ခြင်း။",
            "စွမ်းအင်လျှော့ / အသက်ရှု၊ ဆက်တိုက် လေ့ကျင့်မှု၊ ဆုတံဆိပ်၊ အဆင့်ဇယား။",
            "AI ဖြင့် အသံထွက်အကဲဖြတ်ခြင်း။",
            "ဒက်ရှ်ဘုတ်မှ CSV / PDF ထုတ်ယူခြင်း။",
            "ကျောင်း/အဖွဲ့အစည်း များစွာ စီမံခန့်ခွဲမှု (multi-tenant)။",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(24)
    add_run(footer, "Learn2Success — လုပ်ဆောင်ချက် လမ်းညွှန်ချက် (မြန်မာ)", size=9, color=SLATE, italic=True)
    add_run(footer, f"\nထုတ်ပေးသည့်ရက် {date.today().isoformat()}", size=9, color=SLATE, italic=True)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
