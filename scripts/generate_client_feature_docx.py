#!/usr/bin/env python3
"""Generate Learn2Success client feature guide (.docx)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Learn2Success-Client-Feature-Guide.docx"

# Brand palette
GOLD = RGBColor(0x78, 0x59, 0x00)
GOLD_DARK = RGBColor(0x6D, 0x51, 0x00)
SKY = RGBColor(0x00, 0x63, 0x99)
GREEN = RGBColor(0x00, 0x6E, 0x1C)
SLATE = RGBColor(0x47, 0x55, 0x69)
BODY = RGBColor(0x23, 0x2C, 0x36)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FILL_GOLD = "FFF9EB"
FILL_SKY = "E3F2FD"
FILL_GREEN = "F3FFF2"
FILL_HEADER = "785900"


def set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def style_paragraph(paragraph, *, size=11, bold=False, color=BODY, space_after=6):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = color


def add_run(paragraph, text, *, size=11, bold=False, color=BODY, italic=False):
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    return run


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p, FILL_SKY)
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Learn2Success", size=34, bold=True, color=GOLD)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p2, FILL_SKY)
    p2.paragraph_format.space_after = Pt(4)
    add_run(p2, "Client Feature Guide", size=18, bold=True, color=SKY)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_shading(p3, FILL_SKY)
    p3.paragraph_format.space_after = Pt(24)
    add_run(p3, "Together we learn, together we succeed.", size=12, italic=True, color=GOLD_DARK)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(24)
    add_run(meta, f"Delivered {date.today().strftime('%B %d, %Y')}  ·  Responsive web application", size=10, color=SLATE)


def add_section_heading(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, title, size=16, bold=True, color=GOLD)

    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(10)
        add_run(s, subtitle, size=10, color=SLATE)


def add_subheading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, title, size=12, bold=True, color=SKY)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    style_paragraph(p, size=11, color=BODY, space_after=8)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        add_run(p, item, size=11, color=BODY)


def add_feature_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = table.rows[0].cells
    headers[0].text = "Feature"
    headers[1].text = "What it does for you"
    for cell in headers:
        set_cell_shading(cell, FILL_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(11)
                run.font.name = "Calibri"

    for i, (feature, detail) in enumerate(rows, start=1):
        c0, c1 = table.rows[i].cells
        c0.text = feature
        c1.text = detail
        fill = FILL_GOLD if i % 2 == 1 else "FFFFFF"
        set_cell_shading(c0, fill)
        set_cell_shading(c1, fill)
        for cell in (c0, c1):
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = BODY

    doc.add_paragraph()


def add_role_matrix(doc: Document) -> None:
    headers = ["Capability", "Guest", "Student", "Teacher", "Admin"]
    data = [
        ("View marketing home & brand", "Yes", "—", "—", "—"),
        ("Sign in / sign out", "Yes", "Yes", "Yes", "Yes"),
        ("Week learning journey", "—", "Yes", "—", "—"),
        ("Take quizzes (Choose & Speak)", "—", "Yes", "—", "—"),
        ("Auto-grade (all multiple choice)", "—", "Yes", "—", "—"),
        ("Resume & retake attempts", "—", "Yes", "—", "—"),
        ("View own scores & history", "—", "Yes", "—", "—"),
        ("Staff dashboard with charts", "—", "—", "Yes", "Yes"),
        ("Grade spoken submissions", "—", "—", "Yes", "—"),
        ("Manage users", "—", "—", "—", "Yes"),
        ("Author quizzes", "—", "—", "—", "Yes"),
        ("Manage weeks & quiz order", "—", "—", "—", "Yes"),
    ]

    table = doc.add_table(rows=1 + len(data), cols=5)
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    for i, row in enumerate(data, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            fill = FILL_SKY if i % 2 == 0 else "FFFFFF"
            if j == 0:
                fill = FILL_GOLD if i % 2 == 1 else "FFFFFF"
            set_cell_shading(cell, fill)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if val == "Yes":
                        run.font.bold = True
                        run.font.color.rgb = GREEN
                    else:
                        run.font.color.rgb = BODY

    doc.add_paragraph()


def add_demo_accounts(doc: Document) -> None:
    accounts = [
        ("student", "password", "Week 1 unlocked only — shows locked future weeks"),
        ("student_allweeks", "password", "All 4 weeks unlocked — full journey demo"),
        ("student_jan1", "password", "Registered Jan 1, 2026 — calendar-based unlock demo"),
        ("teacher", "password", "Grading dashboard & submission queue"),
        ("admin", "password", "Full program setup: users, quizzes, weeks"),
    ]

    table = doc.add_table(rows=1 + len(accounts), cols=3)
    table.style = "Table Grid"

    for j, h in enumerate(["Username", "Password", "Best for demo"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)

    for i, (user, pwd, note) in enumerate(accounts, start=1):
        row = table.rows[i].cells
        row[0].text = user
        row[1].text = pwd
        row[2].text = note
        fill = FILL_GREEN if i % 2 == 1 else "FFFFFF"
        for cell in row:
            set_cell_shading(cell, fill)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = BODY

    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title_block(doc)

    add_section_heading(
        doc,
        "1. Platform overview",
        "A single web application for English speaking practice and assessment.",
    )
    add_body(
        doc,
        "Learn2Success helps students practice English through structured quizzes — especially spoken answers "
        "recorded in the browser. Teachers grade submissions with marks and feedback. Administrators run the "
        "program: users, quiz content, and a week-by-week learning journey that unlocks over time.",
    )
    add_body(
        doc,
        "The product is a responsive website (mobile-first for students, desktop-friendly for staff). "
        "No app store install is required.",
    )

    add_section_heading(doc, "2. Who uses the platform")
    add_role_matrix(doc)

    add_section_heading(
        doc,
        "3. Guest — welcome & sign-in",
        "Public pages before login.",
    )
    add_feature_table(
        doc,
        [
            ("Marketing home", "Branded welcome page with logo, tagline, image carousel (swipe & arrows), and feature highlights: Listen, Speak, Succeed."),
            ("Sign in", "Username and password login with optional Remember me."),
            ("Role-based redirect", "After login, each role lands on the right home: student journey, teacher dashboard, or admin dashboard."),
        ],
    )

    add_section_heading(
        doc,
        "4. Student experience",
        "Mobile-friendly learning journey and quiz play.",
    )
    add_subheading(doc, "4.1 Week journey")
    add_bullets(
        doc,
        [
            "Vertical week path (game-style map) on the student home screen.",
            "Weeks unlock on a schedule from the student’s registration date (Week 1 on day 0, Week 2 on day 7, and so on).",
            "Locked weeks stay visible but cannot be opened until unlocked.",
            "Inside each week: numbered quiz steps with Choose / Speak tags and optional time limits.",
        ],
    )
    add_subheading(doc, "4.2 Quizzes & attempts")
    add_bullets(
        doc,
        [
            "Quiz detail shows description, marks, question count, attempt history, and latest status.",
            "Statuses: In progress, Pending (awaiting teacher), Graded (with score).",
            "Start, Continue (resume), and New try — multiple attempts per quiz allowed.",
            "Only one in-progress attempt per quiz at a time.",
        ],
    )
    add_subheading(doc, "4.3 Quiz play")
    add_bullets(
        doc,
        [
            "Choose (multiple choice): select one answer; auto-graded when the whole quiz is multiple choice.",
            "Speak (recording): Listen (text-to-speech), Record with countdown, playback, re-record, then submit audio.",
            "Optional quiz timer — auto-submits when time runs out.",
            "Progress indicator across questions; answers save as the student moves forward.",
        ],
    )

    add_section_heading(
        doc,
        "5. Teacher experience",
        "Focused on grading — not content editing.",
    )
    add_feature_table(
        doc,
        [
            ("Dashboard", "Landing page with KPIs, charts, priority queue (oldest waiting first), and quick links."),
            ("Grading queue", "Filter: Ready to grade, In progress, Graded, All — paginated list."),
            ("By student", "All attempts grouped under each student username."),
            ("Grade screen", "Read-only quiz info; audio playback for Speak answers; per-question marks and feedback."),
            ("Mark validation", "Live total bar; cannot exceed quiz maximum or per-question cap; Save blocked when invalid."),
        ],
    )

    add_section_heading(
        doc,
        "6. Admin experience",
        "Full program and account management.",
    )
    add_feature_table(
        doc,
        [
            ("Dashboard", "Program overview: student/teacher counts, active quizzes, pending grades, charts, recent activity."),
            ("Users", "Create, edit, delete users (student, teacher, admin). Password rules enforced. Cannot delete own admin account."),
            ("Quizzes", "Create and edit quizzes with 1–100 questions. Two types only: Choose and Speak. Optional time limit per quiz."),
            ("Weeks", "Create and edit program weeks. Assign quizzes to a week, drag to reorder, remove from week without deleting quiz."),
            ("Quiz ↔ week rule", "Week assignment lives only in Week edit — not on the quiz form — for a clear admin workflow."),
        ],
    )

    add_section_heading(doc, "7. Question types")
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    for j, h in enumerate(["Type", "Student sees", "Grading"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_shading(cell, FILL_HEADER)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(10)
    rows = [
        ("Choose", "Multiple choice — pick one option", "Automatic when entire quiz is Choose"),
        ("Speak", "Record voice answer in browser", "Teacher grades with marks & feedback"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            set_cell_shading(cell, FILL_GOLD if i == 1 else FILL_SKY)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = BODY
    doc.add_paragraph()

    add_section_heading(
        doc,
        "8. Quality & security",
        "Built for reliable day-to-day use.",
    )
    add_bullets(
        doc,
        [
            "Three roles with strict access control — students cannot open staff areas.",
            "Students can only access their own submissions.",
            "Inactive or week-locked quizzes are blocked.",
            "73 automated tests covering auth, journeys, grading, admin CRUD, and dashboards.",
            "Audio stored securely on server; up to 30 seconds recording per question in the UI.",
        ],
    )

    add_section_heading(
        doc,
        "9. Demo accounts",
        "Use these for walkthroughs and screen recordings.",
    )
    add_demo_accounts(doc)

    add_section_heading(
        doc,
        "10. Not included in this release",
        "Possible future enhancements — not part of the current delivery.",
    )
    add_bullets(
        doc,
        [
            "Native iOS / Android apps (web-only today).",
            "Student self-registration or password reset.",
            "Energy / lives, streaks, badges, or leaderboards.",
            "AI pronunciation scoring.",
            "CSV / PDF export from dashboards.",
            "Multi-school or multi-tenant hierarchy.",
        ],
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(24)
    add_run(footer, "Learn2Success — Client Feature Guide", size=9, color=SLATE, italic=True)
    add_run(footer, f"\nGenerated {date.today().isoformat()}", size=9, color=SLATE, italic=True)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
